import re
import sqlite3
import os
from sqlite3 import Error
from datetime import datetime
import pandas as pd
from icecream import ic
import requests
import time

def create_connection(db_file):
    """创建数据库连接"""
    # 确保数据库文件的目录存在
    db_dir = os.path.dirname(db_file)
    if not os.path.exists(db_dir):
        os.makedirs(db_dir)
    conn = None
    try:
        conn = sqlite3.connect(db_file)
        print("Database connection established.")
    except Error as e:
        print(e)
    return conn

def create_table(conn):
    """创建表"""
    try:
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS articles (
                article_id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                url TEXT UNIQUE NOT NULL,
                content TEXT NOT NULL,
                mp_name TEXT NOT NULL,
                created DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
    except Error as e:
        print(e)

def save_article(conn, article_name, article_url, content, nickname, created):
    """保存文章到数据库"""
    if isinstance(created, str):
        # 如果created是字符串，尝试将其转换为datetime对象
        created = datetime.strptime(created, "%Y-%m-%d %H:%M:%S") if created else None
    try:
        c = conn.cursor()
        c.execute('''
            INSERT OR REPLACE INTO articles(title, url, content, mp_name, created)
            VALUES(?, ?, ?, ?, ?)
        ''', (article_name, article_url, content, nickname, created))
        conn.commit()
        print("Article saved successfully.")
    except Error as e:
        print(e)

def get_url_from_csv(filename):# 指定CSV文件路径
    csv_file_path = 'data\{}'.format(filename)

    # 使用pandas的read_csv函数读取CSV文件
    df = pd.read_csv(csv_file_path)

    # 打印出DataFrame对象，查看数据
    ic(df.head())

    # 假设CSV文件中包含一个名为'url'的列，我们想要获取这个列的所有数据
    url_column = df['url']

    # 打印URL列的所有数据
    ic(url_column)

    # 如果你想要将URLs存储到一个列表中
    urls = url_column.tolist()

    # 打印URL列表
    ic(len(urls))
    return urls

def update_article_content(nickname,csv_name):
    database = "./mp.db"

    # 创建数据库连接
    conn = create_connection(database)

    if conn is not None:
        # 创建表
        create_table(conn)

        # 验证数据
        title = ""
        url = ""
        content = ""
        now = datetime.now()  # 日期字符串
        created = now.strftime("%Y-%m-%d %H:%M:%S")  # 转换为datetime对象
        
        # urls = get_url_from_db(conn, nickname)
        urls = get_url_from_csv(csv_name)
        for url in urls:
            title, content = get_article_content(url)
            # 保存文章
            ic(save_article(conn, title, url, content, nickname, created))
            time.sleep(2)

        # 关闭连接
        conn.close()
    else:
        print("Error! cannot create the database connection.")

def export_articles(nickname):
    dir_path = os.path.join('data', nickname)
    os.makedirs(dir_path, exist_ok=True)
    database = "./mp.db"

    # 创建数据库连接
    conn = create_connection(database)

    if conn is not None:
        """从数据库中读取特定nickname的url记录"""
        c = conn.cursor()
        # 执行SELECT语句，添加WHERE子句来过滤nickname
        c.execute("SELECT title,content FROM articles WHERE mp_name = ?", (nickname,))
        # 获取查询结果
        records = c.fetchall()
        ic(len(records))

        for rec in records[:]:
            filename = rec[0].replace(" ", "").replace("\n","").replace("|","").replace("\\","").replace("/","").replace("<","").replace("?","")
            dir_path = os.path.join('data', nickname)
            file_path = os.path.join(dir_path, filename[:30])
            rec1 = rec[1].replace("\n", "\n- ")
            try:
                with open(f'{file_path}.md', 'w', encoding='utf-8') as f:
                    f.write(f'{rec1}')
                ic(file_path)
            except Error as e:
                ic(file_path,e)

        # 关闭连接
        conn.close()
    else:
        print("Error! cannot create the database connection.")

def get_article_content(url):

    # 要访问的网页链接
    url = 'https://r.jina.ai/'+ url 
    ic(url)

    # 发送GET请求
    response = requests.get(url, verify=False)

    # 检查请求是否成功
    if response.status_code == 200:
        # 获取网页的文本内容
        content = response.text

        # 正则表达式模式，匹配Title:后面的内容
        # \s* 匹配任意数量的空白字符
        # [^:]* 匹配冒号前的所有字符
        # (.*) 捕获冒号及之后的所有字符
        # pattern = r'Title:\s*([^:]*)(?=(\s*$|\n))'
        pattern = r'Title:\s*(.*(?:\n|$))'

        # 使用re.search()查找匹配项
        match = re.search(pattern, content)

        # 如果找到匹配项，则打印出来
        if match:
            print('匹配到的标题:', match.group(1))
            return match.group(1), content
        else:
            print('没有找到匹配的标题')
            return "", content

    else:
        return "", ""
        
# 从数据库中读取特定nickname的url记录
def get_url_from_db(conn, nickname):
    """从数据库中读取特定nickname的url记录"""
    try:
        c = conn.cursor()
        # 执行SELECT语句，添加WHERE子句来过滤nickname
        c.execute("SELECT url FROM articles WHERE mp_name = ? AND content = ''", (nickname,))
        # 获取查询结果
        urls = c.fetchall()
        urls = [url[0] for url in urls] 
        ic(urls)
        return urls
    except Error as e:
        print(e)
        return None

def moc_create(nickname):
    database = "./mp.db"

    # 创建数据库连接
    conn = create_connection(database)

    if conn is not None:
        """从数据库中读取特定nickname的url记录"""
        try:
            c = conn.cursor()
            # 执行SELECT语句，添加WHERE子句来过滤nickname
            c.execute("SELECT title FROM articles WHERE mp_name = ?", (nickname,))
            # 获取查询结果
            records = c.fetchall()
            file_path = os.path.join('data', f'{nickname}',f'moc_{nickname}')
            ic(file_path)
            for rec in records[:]:
                line = rec[0][:30].replace(" ", "").replace("\n","").replace("|","").replace("\\","").replace("/","").replace("<","").replace('?', '')
                if (len(line)):
                    line = f'[[{line}]]\n'
                    with open(f'{file_path}.md', 'a', encoding='utf-8') as f:
                        f.write(line)
                
        except Error as e:
            print(e)

        # 关闭连接
        conn.close()
    else:
        print("Error! cannot create the database connection.")

def export_articles_doc(nickname):
    dir_path = os.path.join('data', nickname)
    os.makedirs(dir_path, exist_ok=True)
    database = "./mp.db"

    conn = create_connection(database)
    if conn is not None:
        try:
            c = conn.cursor()
            c.execute("SELECT title, content FROM articles WHERE mp_name = ?", (nickname,))
            records = c.fetchall()
            
            try:
                from docx import Document
            except ImportError:
                import subprocess
                import sys
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                from docx import Document
            
            for title, content in records:
                filename = re.sub(r'[^\w\-_\. ]', '', title)[:30]
                file_path = os.path.join(dir_path, f"{filename}.docx")
                
                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(content)
                doc.save(file_path)
                
            print(f"Exported {len(records)} articles for {nickname}")
        except Error as e:
            print(f"Database error: {e}")
        finally:
            conn.close()
    else:
        print("Error! Cannot create the database connection.")



if __name__ == '__main__':
    nickname = "L先生说"
    csv_name = nickname +'_p3.csv'
    update_article_content(nickname, csv_name) 
    export_articles(nickname)
    moc_create(nickname)
    export_articles_doc(nickname)

    
    